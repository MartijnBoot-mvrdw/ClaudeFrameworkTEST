#!/usr/bin/env python3
"""
Generate SDD .docx from a JSON data file, using the MvR Word template for styles.

Usage:
    python generate_sdd_docx.py <data.json> <output.docx> [<template.docx>]

Dependencies:
    pip install python-docx

The JSON data file is written by the /solutions Claude Code skill.
The template .docx is used for styles only — its content is replaced entirely.
"""

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed.")
    print("Run: pip install python-docx")
    sys.exit(1)


# ── Style helpers ──────────────────────────────────────────────────────────────

HEADER_BG  = "2E5D9E"   # MvR dark blue
HEADER_FG  = "FFFFFF"


def _shd_element(fill_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    return shd


def set_cell_bg(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(_shd_element(fill_hex))


def style_header_cell(cell, text: str):
    set_cell_bg(cell, HEADER_BG)
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_table(doc, headers: list[str], rows: list[list], style="Table Grid"):
    """Create a table with a styled header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = style
    except KeyError:
        table.style = "Table Grid"
    for i, h in enumerate(headers):
        style_header_cell(table.rows[0].cells[i], h)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val) if val is not None else ""
    return table


def add_kv_table(doc, pairs: list[tuple], style="Table Grid"):
    """Two-column key/value table with a header row."""
    return add_table(doc, ["Eigenschap", "Waarde"], pairs, style)


def gap(doc):
    doc.add_paragraph()


# ── Document builder ───────────────────────────────────────────────────────────

def _clear_body(doc):
    """Remove all content from body, preserving section properties."""
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def build(data: dict, output_path: str, template_path: str | None = None):
    # Load template for styles; clear its content so we start fresh.
    if template_path and Path(template_path).exists():
        doc = Document(template_path)
        _clear_body(doc)
    else:
        doc = Document()
        s = doc.sections[0]
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3.0)
        s.right_margin  = Cm(2.5)

    v = lambda key, fallback="<<ONBEKEND>>": data.get(key) or fallback

    # ── Title block ────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    r = t.add_run("Solution Design Document")
    r.bold = True
    r.font.size = Pt(24)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.add_run(v("procesnaam")).font.size = Pt(16)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    gap(doc)
    meta = doc.add_table(rows=3, cols=2)
    try:
        meta.style = "Table Grid"
    except KeyError:
        pass
    for i, (k, val) in enumerate([
        ("Project", v("projectnaam")),
        ("Afdeling", v("afdeling")),
        ("Datum",    v("datum")),
    ]):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = val

    doc.add_page_break()

    # ── Versie- en revisietabel ────────────────────────────────────────────────
    doc.add_heading("Versie- en revisietabel", level=2)
    versies = data.get("versie_tabel") or [{
        "versie": "0.1",
        "datum":  v("datum"),
        "auteur": "<<AUTEUR>>",
        "omschrijving": "Eerste concept op basis van PDD",
    }]
    add_table(
        doc,
        ["Versie", "Datum", "Auteur", "Omschrijving"],
        [[r.get("versie",""), r.get("datum",""), r.get("auteur",""), r.get("omschrijving","")] for r in versies],
    )
    gap(doc)

    # ── Stakeholders ───────────────────────────────────────────────────────────
    doc.add_heading("Stakeholders", level=2)
    add_table(
        doc,
        ["Rol", "Naam (Klant)", "Naam (MvR DW)"],
        [[s.get("rol",""), s.get("naam_klant","<<ONBEKEND>>"), s.get("naam_mvr","<<ONBEKEND>>")] for s in data.get("stakeholders", [])],
    )
    gap(doc)
    doc.add_page_break()

    # ══ SECTIE 1 ══════════════════════════════════════════════════════════════
    doc.add_heading("1 Overzicht oplossing", level=1)

    # 1.1
    doc.add_heading("1.1 Samenvatting", level=2)
    add_kv_table(doc, [
        ("Procesnaam",                    v("procesnaam")),
        ("Projectnaam",                   v("projectnaam")),
        ("Afdeling",                      v("afdeling")),
        ("Beschrijving",                  v("beschrijving")),
        ("Frequentie",                    v("frequentie")),
        ("Verwacht aantal transacties",   v("verwacht_aantal_transacties")),
        ("Maximale doorlooptijd",         v("maximale_doorlooptijd")),
        ("Type robot",                    v("type_robot")),
        ("Transactietype",                v("transactietype")),
        ("Dispatcher/Performer",          v("dispatcher_performer")),
    ])
    gap(doc)

    # 1.2
    doc.add_heading("1.2 Procesflow", level=2)
    add_table(
        doc,
        ["Nr.", "Sub-proces", "Applicatie", "Handmatig / Geautomatiseerd", "Voorganger"],
        [[str(p.get("nr","")), p.get("sub_proces",""), p.get("applicatie",""),
          p.get("type",""), str(p.get("voorganger",""))]
         for p in data.get("procesflow", [])],
    )
    gap(doc)

    # 1.3
    doc.add_heading("1.3 Decompositie Processtappen", level=2)
    if data.get("decompositie_prose"):
        doc.add_paragraph(data["decompositie_prose"])
    for stap in data.get("decompositie_stappen", []):
        doc.add_paragraph(str(stap), style="List Paragraph")
    if data.get("uitvalpad"):
        p = doc.add_paragraph()
        p.add_run("Uitvalpad: ").bold = True
        p.add_run(data["uitvalpad"])
    gap(doc)

    # 1.4
    doc.add_heading("1.4 Benodigde rechten, applicaties en functionaliteiten", level=2)
    add_table(
        doc,
        ["Applicatie", "Type", "Browser", "Rechten benodigd", "Opmerkingen"],
        [[r.get("applicatie",""), r.get("type",""), r.get("browser",""),
          r.get("rechten",""), r.get("opmerkingen","")]
         for r in data.get("rechten", [])],
    )
    gap(doc)
    doc.add_heading("1.4.1 Browser", level=3)
    doc.add_paragraph(v("browser_settings"))
    gap(doc)

    # ══ SECTIE 2 ══════════════════════════════════════════════════════════════
    doc.add_heading("2 Beschrijving technische workflow", level=1)

    # 2.1
    doc.add_heading("2.1 Algemeen robotontwerp", level=2)
    if data.get("robotontwerp_prose"):
        doc.add_paragraph(data["robotontwerp_prose"])
    add_kv_table(doc, [
        ("Framework",           "MvR_REFramework"),
        ("Dispatcher aanwezig", v("dispatcher_aanwezig", "Nee")),
        ("Queue naam",          v("queue_naam", data.get("projectnaam","<<ONBEKEND>>"))),
        ("QueueRetry",          v("queue_retry", "Nee")),
        ("MaxRetry (framework)", v("max_retry", "0")),
    ])
    gap(doc)

    # 2.2
    doc.add_heading("2.2 Queue en retry mechanisme", level=2)
    if data.get("retry_prose"):
        doc.add_paragraph(data["retry_prose"])
    add_table(
        doc,
        ["Type uitval", "Trigger", "Gevolg"],
        [[u.get("type",""), u.get("trigger",""), u.get("gevolg","")]
         for u in data.get("uitval_tabel", [])],
    )
    gap(doc)

    # 2.3
    doc.add_heading("2.3 Init-fase", level=2)
    if data.get("init_prose"):
        doc.add_paragraph(data["init_prose"])
    add_table(
        doc,
        ["Asset naam", "Type", "Omschrijving", "Waarde (indien bekend)"],
        [[a.get("naam",""), a.get("type",""), a.get("omschrijving",""), a.get("waarde","")]
         for a in data.get("assets", [])],
    )
    gap(doc)

    # 2.4 – 2.6
    for heading, key in [
        ("2.4 GetTransactionData-fase", "get_transaction_prose"),
        ("2.5 Procesfase",              "proces_prose"),
        ("2.6 Eindprocesfase",          "eindproces_prose"),
    ]:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(v(key))
        gap(doc)

    # ══ SECTIE 3 ══════════════════════════════════════════════════════════════
    doc.add_heading("3 Omgevingsafhankelijkheden", level=1)
    if data.get("omgeving_prose"):
        doc.add_paragraph(data["omgeving_prose"])
    if data.get("omgevingen"):
        add_table(
            doc,
            ["Omgeving", "Eigenschap", "Waarde", "Opmerkingen"],
            [[o.get("omgeving",""), o.get("eigenschap",""), o.get("waarde",""), o.get("opmerkingen","")]
             for o in data["omgevingen"]],
        )
    gap(doc)

    # ══ SECTIE 4 ══════════════════════════════════════════════════════════════
    doc.add_heading("4 Processtappen", level=1)

    for stap in data.get("processtappen", []):
        nr  = stap.get("nr", "?")
        sub = stap.get("sub_proces", "<<ONBEKEND>>")
        app = stap.get("applicatie", "<<ONBEKEND>>")

        doc.add_heading(f"4.{nr} Stap {nr}: {sub} ({app})", level=2)

        doc.add_heading(f"4.{nr}.1 Algemene informatie", level=3)
        add_kv_table(doc, [
            ("Workflowbestand",   stap.get("workflowbestand", "<<ONBEKEND>>")),
            ("Applicatie",        app),
            ("Doel",              stap.get("doel",              "<<ONBEKEND>>")),
            ("Navigatie",         stap.get("navigatie",         "<<ONBEKEND>>")),
            ("Input argumenten",  stap.get("input_argumenten",  "<<ONBEKEND>>")),
            ("Output argumenten", stap.get("output_argumenten", "(geen)")),
        ])
        gap(doc)

        doc.add_heading(f"4.{nr}.2 Omschrijving handelingen", level=3)
        for handeling in stap.get("handelingen", []):
            doc.add_paragraph(str(handeling), style="List Paragraph")
        gap(doc)

    # ── Save ───────────────────────────────────────────────────────────────────
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"SDD .docx gegenereerd: {output_path}")


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python generate_sdd_docx.py <data.json> <output.docx> [<template.docx>]")
        sys.exit(1)

    data_file     = sys.argv[1]
    output_file   = sys.argv[2]
    template_file = sys.argv[3] if len(sys.argv) > 3 else None

    with open(data_file, encoding="utf-8") as f:
        sdd_data = json.load(f)

    build(sdd_data, output_file, template_file)
